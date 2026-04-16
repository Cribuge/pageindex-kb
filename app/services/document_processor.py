"""
Document text extraction service (no chunking - preserves full text for PageIndex).
"""
from typing import List, Optional, Tuple
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
import pandas as pd


class DocumentProcessor:

    def extract_text_from_pdf(self, file_path: str) -> List[Tuple[str, dict]]:
        chunks_with_meta = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    chunks_with_meta.append((
                        text,
                        {"page": page_num + 1, "source": "pdf"}
                    ))
        return chunks_with_meta

    def extract_text_from_docx(self, file_path: str) -> List[Tuple[str, dict]]:
        doc = DocxDocument(file_path)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        full_text = "\n\n".join(paragraphs)

        tables = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    tables.append(row_text)
            tables.append("")

        if tables:
            full_text += "\n\n[Tables]\n" + "\n".join(tables)

        if full_text.strip():
            return [(full_text, {"source": "docx"})]
        return []

    def extract_text_from_txt(self, file_path: str) -> List[Tuple[str, dict]]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        return [(text, {"source": "txt"})] if text.strip() else []

    def extract_text_from_excel(self, file_path: str) -> List[Tuple[str, dict]]:
        chunks_with_meta = []
        excel_file = pd.ExcelFile(file_path)
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            text = f"[Sheet: {sheet_name}]\n" + df.to_string(index=False)
            if text.strip():
                chunks_with_meta.append((text, {"source": "excel", "sheet": sheet_name}))
        return chunks_with_meta

    def process_file(self, file_path: str, file_type: Optional[str] = None) -> List[Tuple[str, dict]]:
        if file_type is None:
            file_type = Path(file_path).suffix.lower().lstrip(".")

        extractors = {
            "pdf": self.extract_text_from_pdf,
            "docx": self.extract_text_from_docx,
            "doc": self.extract_text_from_docx,
            "txt": self.extract_text_from_txt,
            "md": self.extract_text_from_txt,
            "xlsx": self.extract_text_from_excel,
            "xls": self.extract_text_from_excel,
        }

        extractor = extractors.get(file_type)
        if not extractor:
            return []

        try:
            return extractor(file_path)
        except Exception as e:
            print(f"Error processing {file_path}: {e}")
            return []


document_processor = DocumentProcessor()
